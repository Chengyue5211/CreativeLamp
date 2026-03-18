"""
将软著登记材料从 Markdown/TXT 转换为规范的 Word (.docx) 文档
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

LEGAL_DIR = Path(__file__).parent


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            from lxml import etree
            element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))


def create_base_doc(title_text, margin_cm=2.54):
    """创建基础文档，设置页边距和默认字体"""
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题样式
    for i in range(1, 5):
        heading = doc.styles[f'Heading {i}']
        heading.font.name = '黑体'
        heading.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            heading.font.size = Pt(22)
        elif i == 2:
            heading.font.size = Pt(16)
        elif i == 3:
            heading.font.size = Pt(14)
        else:
            heading.font.size = Pt(12)

    return doc


def add_cover_page(doc, title, subtitle="", version="V1.0", date="2026年3月"):
    """添加封面页"""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    if subtitle:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for _ in range(4):
        doc.add_paragraph("")

    info_lines = [
        f"版本号：{version}",
        f"编制日期：{date}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


def md_to_docx(md_path, output_path, cover_title, cover_subtitle=""):
    """将 Markdown 转换为 Word 文档"""
    doc = create_base_doc(cover_title)
    add_cover_page(doc, cover_title, cover_subtitle)

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                pf = p.paragraph_format
                pf.left_indent = Cm(1)
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)
                # 灰色背景通过shading
                from docx.oxml.ns import qn as _qn
                from lxml import etree
                shading = etree.SubElement(p._p.get_or_add_pPr(), _qn('w:shd'))
                shading.set(_qn('w:fill'), 'F5F5F5')
                shading.set(_qn('w:val'), 'clear')
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            # 跳过分隔行
            if all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # 检查下一行是否还是表格
            if i + 1 < len(lines) and '|' in lines[i + 1] and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # 输出表格
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < num_cols:
                                cell = table.cell(ri, ci)
                                # 清除默认段落
                                cell.paragraphs[0].text = ""
                                run = cell.paragraphs[0].add_run(
                                    re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text.replace('<br>', '\n'))
                                )
                                run.font.size = Pt(9)
                                run.font.name = '宋体'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                if ri == 0:
                                    run.bold = True

                    doc.add_paragraph("")  # 表后空行
                    table_rows = []
                i += 1
                continue

        # 水平线
        if line.strip() == '---':
            i += 1
            continue

        # 标题
        if line.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.+)', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                doc.add_heading(text, level=level)
                i += 1
                continue

        # 引用块
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # 有序/无序列表
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1 + indent * 0.5)
            # 处理加粗
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            bullet = '- ' if list_match.group(2) in ['-', '*'] else f"{list_match.group(2)} "
            first = True
            for part in parts:
                bold_match = re.match(r'\*\*([^*]+)\*\*', part)
                if first:
                    r = p.add_run(bullet)
                    r.font.size = Pt(10.5)
                    r.font.name = '宋体'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    first = False
                if bold_match:
                    r = p.add_run(bold_match.group(1))
                    r.bold = True
                    r.font.size = Pt(10.5)
                    r.font.name = '宋体'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    r = p.add_run(part)
                    r.font.size = Pt(10.5)
                    r.font.name = '宋体'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        text = line.strip()
        p = doc.add_paragraph()
        # 处理加粗
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            bold_match = re.match(r'\*\*([^*]+)\*\*', part)
            if bold_match:
                r = p.add_run(bold_match.group(1))
                r.bold = True
            else:
                r = p.add_run(part)
            r.font.size = Pt(10.5)
            r.font.name = '宋体'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        i += 1

    doc.save(str(output_path))
    print(f"  [OK] {output_path.name}")


def source_code_to_docx(txt_path, output_path, page_label):
    """将源代码 TXT 转换为软著登记格式的 Word 文档（每页50行）"""
    doc = create_base_doc("源代码", margin_cm=2.0)

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    # 添加页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"绘创前程 儿童原创视觉表达成长平台软件 V1.0 — 源代码（{page_label}）")
    run.font.size = Pt(8)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(128, 128, 128)

    content = txt_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    # 每页50行
    page_size = 50
    for page_start in range(0, len(lines), page_size):
        page_lines = lines[page_start:page_start + page_size]
        code_text = '\n'.join(page_lines)

        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        p.paragraph_format.line_spacing = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 不是最后一页则分页
        if page_start + page_size < len(lines):
            doc.add_page_break()

    doc.save(str(output_path))
    print(f"  [OK] {output_path.name}")


def main():
    print("正在生成 Word 文档...\n")

    # 1. 软著操作手册
    md_to_docx(
        LEGAL_DIR / '软著操作手册.md',
        LEGAL_DIR / '软著操作手册.docx',
        '绘创前程',
        '软件操作手册'
    )

    # 2. 知识产权保护与专利分析
    md_to_docx(
        LEGAL_DIR / '知识产权保护与专利分析.md',
        LEGAL_DIR / '知识产权保护与专利分析.docx',
        '绘创前程',
        '知识产权保护与专利分析'
    )

    # 3. 源代码文档（前30页）
    source_code_to_docx(
        LEGAL_DIR / '软著源代码_前30页.txt',
        LEGAL_DIR / '软著源代码_前30页.docx',
        '前30页'
    )

    # 4. 源代码文档（后30页）
    source_code_to_docx(
        LEGAL_DIR / '软著源代码_后30页.txt',
        LEGAL_DIR / '软著源代码_后30页.docx',
        '后30页'
    )

    print("\n全部完成！文件位于 D:\\CreativeLamp\\docs\\legal\\")
    print("  - 软著操作手册.docx")
    print("  - 知识产权保护与专利分析.docx")
    print("  - 软著源代码_前30页.docx")
    print("  - 软著源代码_后30页.docx")


if __name__ == '__main__':
    main()
