"""生成《绘创前程双训练系统简版教程》Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.font.color.rgb = RGBColor(0x4E, 0x8D, 0x7C)
    hs._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("绘创前程")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x4E, 0x8D, 0x7C)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("双训练系统使用教程")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("没有原型，就没有变形；没有变形，就没有创造")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("儿童原创视觉表达成长平台\n简版教程 · 内部使用")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
doc.add_heading("目录", level=1)
toc_items = [
    "一、平台简介与教育理念",
    "二、模块A — 原型组合创意绘画",
    "三、模块B — 图形变形与结构生成",
    "四、九级年龄梯度体系",
    "五、完整使用流程",
    "六、每日训练任务说明",
    "七、常见问题",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 一、平台简介
# ============================================================
doc.add_heading("一、平台简介与教育理念", level=1)

doc.add_heading("1.1 平台定位", level=2)
doc.add_paragraph(
    "「绘创前程」是一个面向4-12岁儿童的原创视觉表达成长平台。"
    "平台采用「软件示范 → 纸上创作 → 拍照上传 → 评估归档」的闭环模式，"
    "让孩子在真实纸面上完成创作，而非在屏幕上涂鸦。"
)

doc.add_heading("1.2 核心教育理念", level=2)
p = doc.add_paragraph()
run = p.add_run("核心公式：没有原型，就没有变形；没有变形，就没有创造。")
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    "这意味着：孩子需要先掌握基础的视觉「原型」（线条、图形），"
    "再学会对它们进行「变形」（拉长、切分、叠加等），"
    "最终在变形的基础上实现真正的「创造」。"
    "这就是为什么平台设计了两个训练模块——"
    "模块A负责「原型」积累，模块B负责「变形」训练。"
)

doc.add_heading("1.3 双训练体系概览", level=2)

# 对比表格
table = doc.add_table(rows=6, cols=3, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["", "模块A：原型组合创意绘画", "模块B：图形变形与结构生成"]
data = [
    ["培养目标", "积累视觉语言词汇，学会用线条装饰和表达", "掌握图形变形能力，学会从简单图形生成复杂造型"],
    ["核心素材", "108种线条原型（9大类×12变式）", "9种基础图形 + 7种变形动作"],
    ["创作方式", "用不同线条原型装饰轮廓图或自由创作", "对基础图形施加变形动作，生成新造型"],
    ["输出成果", "装饰过的轮廓画、线条创意画", "奇想生物、梦幻建筑、变形植物等"],
    ["核心能力", "观察力、线条控制力、装饰感", "空间想象力、结构思维、创造力"],
]
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for i, row_data in enumerate(data):
    for j, text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = text

doc.add_page_break()

# ============================================================
# 二、模块A
# ============================================================
doc.add_heading("二、模块A — 原型组合创意绘画", level=1)

doc.add_heading("2.1 什么是「原型」？", level=2)
doc.add_paragraph(
    "原型是最基础的视觉元素——线条。就像语言有字母、音乐有音符，"
    "绘画的「字母」就是各种线条。模块A帮助孩子认识和掌握108种线条原型，"
    "这些线条分为9大类，每类有12种变式。"
)

doc.add_heading("2.2 九大线条类别", level=2)

line_categories = [
    ("直线（Straight）", "最基础的线条，包含水平线、垂直线、对角线、十字线、放射线、网格线等12种变式。是所有线条的基础。"),
    ("虚线（Dashed）", "断续的线条，包含短虚线、长虚线、点线、点划线、粗虚线、渐变虚线等。训练节奏感和韵律感。"),
    ("弧线（Arc）", "弯曲的优美线条，包含彩虹弧、花瓣弧、山形弧、连续弧等。训练手腕灵活度和曲线控制力。"),
    ("波浪线（Wave）", "起伏的线条，包含基础波浪、鱼鳞波、河流波、双波浪等。表现流动感和动态美。"),
    ("折线/锯齿线（Zigzag）", "尖锐转折的线条，包含基础折线、M形、W形、阶梯形等。训练方向转换能力。"),
    ("螺旋线（Spiral）", "旋转的线条，包含圆螺旋、三角螺旋、花形螺旋、蜗牛螺旋等。训练空间旋转感。"),
    ("弹簧线（Spring）", "连续圈圈，包含圆弹簧、椭圆弹簧、渐变弹簧、花形弹簧等。训练连续手绘能力。"),
    ("城墙线（Castle）", "像城墙垛口的线条，包含方城墙、圆顶、尖顶、阶梯、密排等。训练重复和结构感。"),
    ("中式纹样线（Chinese Pattern）", "中国传统装饰纹样，包含回纹、云纹、如意纹、万字纹、卷草纹等。传承文化美学。"),
]

for name, desc in line_categories:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(f"：{desc}")

doc.add_heading("2.3 轮廓图库", level=2)
doc.add_paragraph(
    "平台内置了40多种手绘轮廓图，孩子可以打印出来，用线条原型进行装饰。轮廓图分为6大类别："
)

contour_categories = [
    ("动物类（15种）", "猫、蝴蝶、大象、熊猫、猫头鹰、鲸鱼、螃蟹、变色龙、瓢虫等"),
    ("植物类（6种）", "花朵、树木、叶子、蘑菇、橡果、石榴等"),
    ("物品类（11种）", "花瓶、杯子、雨伞、包包、帽子、瓶子、盘子、茶壶等"),
    ("服饰类（6种）", "各式旗袍、裙子等中国传统服饰"),
    ("交通类（5种）", "汽车、轮船、火箭、潜水艇、热气球"),
    ("幻想类（1种）", "小怪兽——激发自由想象"),
]
for name, items in contour_categories:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(f"：{items}")

doc.add_heading("2.4 组合方式", level=2)
doc.add_paragraph("线条不是单独使用的，孩子需要学会用不同方式将线条组合到画面中：")

composition_modes = [
    "重复：同一种线条反复使用，填满区域",
    "排列：不同线条有序排列，形成图案",
    "疏密变化：有的地方线条密集，有的稀疏",
    "粗细变化：线条从粗到细或从细到粗",
    "方向变化：同一线条朝不同方向排列",
    "环绕：线条围绕中心向外扩展",
    "交叉：不同方向的线条互相交叉",
    "分区填充：轮廓内分成不同区域，每区用不同线条",
    "中心扩散：从中心向四周逐渐变化",
]
for mode in composition_modes:
    doc.add_paragraph(mode, style="List Bullet")

doc.add_heading("2.5 模块A任务类型", level=2)

task_types_a = [
    ("识别任务", "系统展示线条，孩子认出它属于哪一类。适合初学者熟悉线条。"),
    ("模仿任务", "照着示范画出相同的线条。练习手部精细控制。"),
    ("补全任务", "给出一半的图案，孩子补全另一半。训练对称和延续能力。"),
    ("迁移任务", "把学过的线条应用到新的场景中。训练灵活运用。"),
    ("创作任务", "自由选择线条进行创作。最高阶的综合训练。"),
    ("轮廓装饰", "打印轮廓图，用各种线条装饰。最常见的练习方式。"),
    ("扩图任务", "把小草图扩展成完整大画。训练构图和画面扩展能力。"),
]
for name, desc in task_types_a:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}：")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 三、模块B
# ============================================================
doc.add_heading("三、模块B — 图形变形与结构生成", level=1)

doc.add_heading("3.1 什么是「变形」？", level=2)
doc.add_paragraph(
    "变形是创造的核心能力。模块B教孩子从最简单的基础图形出发，"
    "通过一系列变形动作（拉、切、叠、拼、挖、延、添），"
    "把简单形状变成复杂的、有创意的造型——比如把一个圆变成一只猫头鹰，"
    "把三角形变成一座城堡。"
)

doc.add_heading("3.2 九种基础图形", level=2)
doc.add_paragraph("所有变形都从这9种基础图形开始：")

shapes = [
    "圆形 — 最基础最自然的形状",
    "方形 — 稳定、建筑感",
    "三角形 — 尖锐、动态",
    "椭圆形 — 圆的变体，更灵活",
    "半圆形 — 可以做帽子、拱门、彩虹",
    "水滴形 — 自然界常见，如叶子、火焰",
    "叶形 — 植物、翅膀、眼睛的基础",
    "云形 — 柔软、梦幻的造型",
    "胶囊形 — 圆角矩形，现代感",
]
for s in shapes:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.3 七种变形动作", level=2)
doc.add_paragraph("这是模块B的核心——七种改变图形的方法：")

transforms = [
    ("拉长", "改变图形的比例",
     "把圆形纵向拉长变成蛋形，把方形横向压扁变成砖块。",
     "训练比例感知能力"),
    ("切分", "把一个图形切成多个部分",
     "把圆形一分为二变成两个半圆，把方形切成四格变成窗户。",
     "训练结构分析能力"),
    ("叠加", "把多个图形叠在一起",
     "圆形叠在三角形上变成冰淇淋，两个圆叠加变成雪人。",
     "训练层次感和空间关系"),
    ("拼接", "把图形连接在一起",
     "三角形拼在方形上变成房子，圆形拼在长条上变成棒棒糖。",
     "训练组合构建能力"),
    ("缺口与添加", "在图形上挖出缺口或添加突起",
     "圆形挖个缺口变成吃豆人，方形加突起变成机器人。",
     "训练正负空间意识"),
    ("延展", "让图形从边缘向外生长",
     "圆形长出四条腿变成乌龟，方形长出烟囱变成工厂。",
     "训练发散思维和扩展能力"),
    ("增加", "在图形内部添加细节",
     "圆形里加眼睛嘴巴变成表情，方形里加窗户门变成楼房。",
     "训练细节观察和丰富表达"),
]

for name, definition, example, ability in transforms:
    p = doc.add_paragraph()
    run = p.add_run(f"【{name}】")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4E, 0x8D, 0x7C)
    doc.add_paragraph(f"定义：{definition}")
    doc.add_paragraph(f"举例：{example}")
    p2 = doc.add_paragraph(f"能力：{ability}")
    p2.paragraph_format.space_after = Pt(12)

doc.add_heading("3.4 「增加」动作的五个子类型", level=2)
doc.add_paragraph("「增加」是最丰富的变形动作，细分为五种：")

increase_types = [
    ("细节增加", "添加花纹、表情、斑点等视觉细节"),
    ("线条增加", "添加胡须、头发、条纹等线条元素"),
    ("部件增加", "添加翅膀、手臂、轮子等功能部件"),
    ("复制增加", "添加窗户、牙齿、花朵等重复元素"),
    ("附属对象增加", "添加帽子、围巾、旁边的小树等附属物"),
]
for name, desc in increase_types:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("3.5 模块B任务类型", level=2)

task_types_b = [
    ("单形变形", "对一个图形施加一种变形。最基础的练习。\n例：「把圆形拉长，变成你能想到的东西」"),
    ("双形拼接", "用两个图形拼接出新造型。\n例：「用三角形和方形拼接，创造一座建筑」"),
    ("多形生成", "用多个图形和多种变形，创造复杂造型。\n例：「用圆形、三角形和叠加、延展，创造一个小动物」"),
    ("主题生成", "围绕一个主题自由创作。\n例：「主题是海洋，用学过的图形和变形创造海洋生物」"),
    ("系列生成", "创作一组相关的作品。\n例：「创造一个图形家族，每个成员由同一个基础图形变化而来」"),
]
for name, desc in task_types_b:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("3.6 模块B的七种输出类型", level=2)
doc.add_paragraph("通过变形训练，孩子可以创造出这些类型的作品：")

output_types = [
    "奇想生物 — 用变形创造的想象中的动物或角色",
    "梦幻建筑 — 用图形搭建的梦幻房屋和城市",
    "结构型角色 — 有结构感的人物或机器人",
    "变形植物 — 从图形变化出来的奇异植物",
    "生成式物体 — 日常物品的创意变体",
    "主题系列 — 围绕主题创作的系列作品",
    "IP雏形 — 可以发展为个人IP的原创角色",
]
for t in output_types:
    doc.add_paragraph(t, style="List Bullet")

doc.add_page_break()

# ============================================================
# 四、九级年龄梯度
# ============================================================
doc.add_heading("四、九级年龄梯度体系", level=1)

doc.add_paragraph(
    "系统根据孩子的年龄自动分配训练等级，"
    "不同等级对应不同的任务难度和类型。随着等级提升，"
    "任务中使用的原型数量、变形动作数量、以及任务复杂度都会递增。"
)

# 九级表格
table = doc.add_table(rows=10, cols=5, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["等级", "年龄", "模块A特点", "模块B特点", "核心任务"]
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

levels = [
    ("准备级", "3-4岁", "1种线条，认识为主", "1种图形+1种动作", "识别、模仿、单形变形"),
    ("初级上", "4-5岁", "2种线条，开始模仿", "1-2种图形+2种动作", "模仿、补全"),
    ("初级下", "5-6岁", "3种线条，尝试迁移", "1-3种图形+3种动作", "迁移、双形拼接"),
    ("中级上", "6-8岁", "4种线条，开始创作", "1-4种图形+4种动作", "创作、轮廓装饰、增加细节"),
    ("中级下", "8-10岁", "5种线条，轮廓装饰", "1-5种图形+5种动作", "轮廓装饰、多形生成"),
    ("高级上", "10-12岁", "6种线条，扩图创作", "1-6种图形+6种动作", "扩图、主题生成"),
    ("高级下", "12岁+", "6种线条，综合创作", "1-6种图形+7种动作", "主题生成、系列生成"),
    ("超级上", "初中", "7种线条，高阶创作", "1-7种图形+7种动作", "综合创作"),
    ("超级下", "初中+", "8种线条，全面掌握", "1-8种图形+7种动作", "自由创作、系列生成"),
]
for i, (lv, age, a_desc, b_desc, tasks) in enumerate(levels):
    table.rows[i + 1].cells[0].text = lv
    table.rows[i + 1].cells[1].text = age
    table.rows[i + 1].cells[2].text = a_desc
    table.rows[i + 1].cells[3].text = b_desc
    table.rows[i + 1].cells[4].text = tasks

doc.add_page_break()

# ============================================================
# 五、完整使用流程
# ============================================================
doc.add_heading("五、完整使用流程", level=1)

doc.add_heading("第一步：家长注册登录", level=2)
steps_1 = [
    "打开平台网址，进入登录页面",
    "如果没有账号，点击「还没有账号？立即注册」",
    "输入手机号（11位）和密码（至少6位）",
    "如果有朋友的邀请码，可以填写获得30学币奖励",
    "点击「注册」，成功后自动进入家长中心",
]
for i, s in enumerate(steps_1, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第二步：添加孩子", level=2)
steps_2 = [
    "在家长中心点击「+ 添加孩子」按钮",
    "输入孩子的昵称、年龄、性别",
    "系统根据年龄自动分配训练等级（如6岁 → 初级下）",
    "点击「确认添加」，孩子信息出现在首页",
]
for i, s in enumerate(steps_2, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第三步：进入孩子的创作空间", level=2)
steps_3 = [
    "在家长中心点击孩子卡片上的「进入创作空间」按钮",
    "系统切换到孩子模式，显示孩子首页",
    "首页展示：今日训练任务、作品数量、资源库入口",
]
for i, s in enumerate(steps_3, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第四步：生成并查看训练任务", level=2)
steps_4 = [
    "点击「生成今日训练」按钮",
    "系统自动生成2个任务：1个模块A任务 + 1个模块B任务",
    "每个任务卡片会显示：任务类型、使用的线条/图形、预计时间",
    "点击任务卡片进入任务详情页",
]
for i, s in enumerate(steps_4, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第五步：纸上创作", level=2)
p = doc.add_paragraph()
run = p.add_run("这是最重要的环节——孩子在纸上进行真实创作！")
run.bold = True

steps_5 = [
    "查看任务详情中的线条示范/图形示范",
    "如果是轮廓装饰任务，点击「打印底稿」将轮廓图打印到纸上",
    "如果是自由创作，准备一张白纸",
    "孩子参照屏幕上的示范，在纸上进行真实绘画",
    "完成后拍照准备上传",
]
for i, s in enumerate(steps_5, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第六步：拍照上传", level=2)
steps_6 = [
    "在任务详情页点击「上传作品」",
    "拍照或从相册选择作品照片（支持JPG/PNG/WebP，最大10MB）",
    "给作品起一个名字（必填）",
    "可选填作品描述",
    "点击「上传」，作品自动关联到当前任务",
]
for i, s in enumerate(steps_6, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("第七步：评价与成长记录", level=2)
steps_7 = [
    "上传完成后，可以对作品进行评价",
    "评价维度（0-10分）：原创性、细节丰富度、构图表现、创意表达",
    "可以写文字点评",
    "家长也可以在家长中心查看孩子作品并给出评价",
    "所有作品和评价自动记录在成长档案中",
]
for i, s in enumerate(steps_7, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_page_break()

# ============================================================
# 六、每日训练任务说明
# ============================================================
doc.add_heading("六、每日训练任务说明", level=1)

doc.add_heading("6.1 任务生成规则", level=2)
rules = [
    "每天可生成2个训练任务（1个模块A + 1个模块B）",
    "同一天重复点击「生成任务」不会产生新任务，而是显示当天已有的任务",
    "任务内容根据孩子的等级自动调整难度",
    "线条原型和基础图形是随机抽取的，保证每天不重复",
]
for r in rules:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("6.2 任务状态流转", level=2)
doc.add_paragraph("每个任务有4个状态：")

statuses = [
    ("待开始", "任务刚生成，还未开始"),
    ("进行中", "孩子点击了任务，正在创作"),
    ("已提交", "作品已上传，等待评价"),
    ("已评价", "作品已被评价，任务完成"),
]
for name, desc in statuses:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("6.3 模块A任务示例", level=2)
doc.add_paragraph(
    "任务：「轮廓装饰 · 用弧线和波浪线装饰猫咪」\n\n"
    "步骤：\n"
    "1. 系统随机选了「弧线-花瓣弧」和「波浪线-鱼鳞波」两种线条\n"
    "2. 推荐了「猫咪」轮廓图，组合方式为「分区填充」\n"
    "3. 孩子打印猫咪轮廓\n"
    "4. 在猫咪身体不同区域，分别用花瓣弧和鱼鳞波填充\n"
    "5. 拍照上传，获得评价"
)

doc.add_heading("6.4 模块B任务示例", level=2)
doc.add_paragraph(
    "任务：「双形拼接 · 用圆形和三角形创造新造型」\n\n"
    "步骤：\n"
    "1. 系统选了「圆形」和「三角形」两种基础图形\n"
    "2. 变形动作为「拼接」和「增加-细节增加」\n"
    "3. 孩子先把圆形和三角形拼接起来（比如变成冰淇淋、小鸟）\n"
    "4. 再用「增加」动作添加眼睛、花纹等细节\n"
    "5. 拍照上传，获得评价"
)

doc.add_page_break()

# ============================================================
# 七、常见问题
# ============================================================
doc.add_heading("七、常见问题", level=1)

faqs = [
    ("孩子一定要在纸上画吗？",
     "是的。平台的核心理念是「纸屏协同」——屏幕负责示范和出题，"
     "孩子在真实纸面上创作。这样既保护视力，又训练真实的手绘能力。"),
    ("每天必须做两个任务吗？",
     "不强制。两个任务是建议量，孩子可以选择先做一个模块的任务。"
     "重要的是保持创作兴趣，而非机械完成。"),
    ("孩子画得不像怎么办？",
     "「像不像」不是评价标准。平台关注的是：原创性（是否有自己的想法）、"
     "细节（是否丰富）、构图（是否完整）、表达（是否有创意）。"
     "每个孩子都有独特的表达方式。"),
    ("为什么要打印轮廓图？",
     "轮廓图给孩子提供了一个「画框」，让他们专注在线条装饰上，"
     "降低了构图难度。就像填色本一样，但更有创造性——"
     "孩子要自己选择用什么线条来装饰。"),
    ("等级会自动升级吗？",
     "目前等级根据孩子年龄自动确定。未来会加入基于能力评估的升级机制。"),
    ("可以一个家长添加多个孩子吗？",
     "可以。在家长中心点击「+ 添加孩子」即可添加多个孩子，"
     "每个孩子有独立的训练进度和作品档案。"),
    ("如何查看孩子的成长进展？",
     "在家长中心点击孩子卡片下方的「查看全部」可以看到该孩子的所有作品。"
     "孩子模式下也有「成长档案」页面，展示作品统计和成绩趋势。"),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f"Q：{q}")
    run.bold = True
    doc.add_paragraph(f"A：{a}")
    doc.add_paragraph()  # spacing

# ============================================================
# 保存
# ============================================================
output_path = r"D:\CreativeLamp\docs\绘创前程_双训练系统使用教程.docx"
doc.save(output_path)
print(f"教程已生成: {output_path}")
